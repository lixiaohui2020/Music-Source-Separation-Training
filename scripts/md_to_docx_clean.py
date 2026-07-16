#!/usr/bin/env python3
"""Convert Audio Core analysis report Markdown to a clean, well-formatted Word doc."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = Path("/workspace/docs/Audio_Core_分析报告_v2.md")
DST = Path("/workspace/docs/Audio_Core_分析报告_v2.docx")
UPLOAD = Path("/home/ubuntu/.cursor/projects/workspace/uploads/Audio_Core_分析报告_v2.docx")

# Colors
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x5A, 0x88)
GRAY = RGBColor(0x55, 0x55, 0x55)
HEADER_BG = "1F3A5F"
ALT_ROW_BG = "F2F5F8"
NOTE_BG = "FFF8E7"
CODE_BG = "F5F5F5"


def set_run_font(run, size=11, bold=False, italic=False, color=None, name="微软雅黑"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.35, first_line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_rich_runs(paragraph, text: str, base_size=11, base_color=None, bold_default=False):
    """Parse simple **bold** and `code` inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color or NAVY)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 1, bold=False, color=RGBColor(0x8B, 0x00, 0x00), name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, bold=bold_default, color=base_color)


def add_heading(doc, text: str, level: int):
    # Strip markdown heading markers if present
    text = re.sub(r"^#+\s*", "", text).strip()
    p = doc.add_paragraph()
    if level == 0:
        set_paragraph_format(p, space_before=0, space_after=8, line_spacing=1.2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=22, bold=True, color=NAVY)
    elif level == 1:
        set_paragraph_format(p, space_before=18, space_after=8, line_spacing=1.2)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        # bottom border line under H1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1F3A5F")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_paragraph_format(p, space_before=12, space_after=6, line_spacing=1.2)
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=ACCENT)
    else:
        set_paragraph_format(p, space_before=8, space_after=4, line_spacing=1.2)
        run = p.add_run(text)
        set_run_font(run, size=11, bold=True, color=ACCENT)
    return p


def add_body(doc, text: str, indent=0):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    add_rich_runs(p, text, base_size=11)
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=3, line_spacing=1.35)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    bullet = "• " if level == 0 else "– "
    run = p.add_run(bullet)
    set_run_font(run, size=11, color=ACCENT)
    add_rich_runs(p, text, base_size=11)
    return p


def add_numbered(doc, text: str, number: int):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=3, line_spacing=1.35)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    run = p.add_run(f"{number}. ")
    set_run_font(run, size=11, bold=True, color=ACCENT)
    add_rich_runs(p, text, base_size=11)
    return p


def add_note(doc, lines: list[str]):
    """Callout / blockquote style. Accepts a list of lines for multi-line notes."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table, 16.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_BG)
    set_cell_borders(cell, color="E0C080", sz="8")
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    cell.text = ""
    first = True
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        set_paragraph_format(p, space_before=0, space_after=2, line_spacing=1.3)
        add_rich_runs(p, line, base_size=10, base_color=GRAY)
    sp = doc.add_paragraph()
    set_paragraph_format(sp, space_before=0, space_after=6, line_spacing=1.0)
    return table


def add_code_block(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_borders(cell, color="DDDDDD", sz="4")
    set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
    cell.text = ""
    # put each line as a paragraph for readability
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        set_paragraph_format(p, space_before=0, space_after=0, line_spacing=1.15)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=8, name="Consolas", color=RGBColor(0x33, 0x33, 0x33))
    sp = doc.add_paragraph()
    set_paragraph_format(sp, space_before=0, space_after=4, line_spacing=1.0)
    return table


def set_table_full_width(table, width_cm=16.5):
    """Force table to usable page width so columns wrap instead of overflowing."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    # remove existing tblW
    for child in list(tblPr):
        if child.tag == qn("w:tblW"):
            tblPr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))  # approx twips (cm*567)
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # layout fixed
    for child in list(tblPr):
        if child.tag == qn("w:tblLayout"):
            tblPr.remove(child)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def add_table(doc, headers: list[str], rows: list[list[str]]):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_full_width(table, 16.5)

    # font size shrinks for wide tables
    font_size = 9 if ncols <= 5 else (8 if ncols <= 7 else 7)

    # equal column widths
    col_w = int(16.5 * 567 / ncols)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag == qn("w:tcW"):
                    tcPr.remove(child)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_w))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    # header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell, color="1F3A5F", sz="4")
        set_cell_margins(cell, top=30, bottom=30, left=40, right=40)
        p = cell.paragraphs[0]
        set_paragraph_format(p, space_before=0, space_after=0, line_spacing=1.1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_runs(
            p,
            h.strip(),
            base_size=font_size,
            base_color=RGBColor(0xFF, 0xFF, 0xFF),
            bold_default=True,
        )

    # body
    for i, row in enumerate(rows):
        bg = ALT_ROW_BG if i % 2 == 1 else "FFFFFF"
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            cell = table.cell(i + 1, j)
            cell.text = ""
            set_cell_shading(cell, bg)
            set_cell_borders(cell, color="CCCCCC", sz="4")
            set_cell_margins(cell, top=30, bottom=30, left=40, right=40)
            p = cell.paragraphs[0]
            set_paragraph_format(p, space_before=0, space_after=0, line_spacing=1.1)
            add_rich_runs(p, val.strip(), base_size=font_size)

    sp = doc.add_paragraph()
    set_paragraph_format(sp, space_before=0, space_after=8, line_spacing=1.0)
    return table


def add_architecture_diagram(doc):
    """Replace messy ASCII art with a clean nested layout table."""
    # Outer: Audio Core
    outer = doc.add_table(rows=4, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(outer, 16.5)

    # Title row
    c = outer.cell(0, 0)
    c.text = ""
    set_cell_shading(c, HEADER_BG)
    set_cell_borders(c, color="1F3A5F")
    set_cell_margins(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, space_before=0, space_after=0)
    run = p.add_run("Audio Core")
    set_run_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Two engines row — use nested table via cell merge alternative: put text blocks
    c = outer.cell(1, 0)
    c.text = ""
    set_cell_shading(c, "FFFFFF")
    set_cell_borders(c, color="1F3A5F")
    set_cell_margins(c, top=60, bottom=60, left=60, right=60)

    # Build inner 1x2 table manually inside by adding a real table after... 
    # python-docx nesting is awkward; use two paragraphs with clear labels instead.
    blocks = [
        ("AFE Engine（硬化）", "FFT / IFFT · Mel / MFCC · 重采样 · Window / OLA · 激活 LUT"),
        ("NN Execution Unit（可编程）", "1D MAC 阵列（16×32 或 8×64）· 序列控制器 / 微码 · State RegFile（hidden state 本地）"),
        ("共享资源", "L2 SRAM（容量见 §7） · Sequencer / DMA / Ctrl"),
    ]
    first = True
    for title, detail in blocks:
        if first:
            p = c.paragraphs[0]
            first = False
        else:
            p = c.add_paragraph()
        set_paragraph_format(p, space_before=2, space_after=2, line_spacing=1.2)
        run = p.add_run("▸ " + title)
        set_run_font(run, size=10, bold=True, color=NAVY)
        p2 = c.add_paragraph()
        set_paragraph_format(p2, space_before=0, space_after=6, line_spacing=1.2)
        p2.paragraph_format.left_indent = Cm(0.4)
        run = p2.add_run(detail)
        set_run_font(run, size=9, color=GRAY)

    # Interface row
    c = outer.cell(2, 0)
    c.text = ""
    set_cell_shading(c, "E8EEF5")
    set_cell_borders(c, color="1F3A5F")
    set_cell_margins(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, space_before=0, space_after=0)
    run = p.add_run("接口：PDM / I2S / TDM 直入  ·  ↔ AXI / 共享 SRAM / 中断")
    set_run_font(run, size=9, bold=True, color=ACCENT)

    # NPU row
    c = outer.cell(3, 0)
    c.text = ""
    set_cell_shading(c, "F7F7F7")
    set_cell_borders(c, color="888888")
    set_cell_margins(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, space_before=0, space_after=0)
    run = p.add_run("现有 NPU（不动）")
    set_run_font(run, size=10, bold=True, color=GRAY)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p2, space_before=2, space_after=0)
    run = p2.add_run("大 Transformer / MatMul / Conv2D  ·  ASR Conformer、TTS、SepFormer 主体、AST")
    set_run_font(run, size=9, color=GRAY)

    sp = doc.add_paragraph()
    set_paragraph_format(sp, space_before=0, space_after=8, line_spacing=1.0)
    return outer


def add_wakeup_chain(doc):
    """Clean wake-up chain instead of ASCII tree."""
    steps = [
        ("① Mic 输入", "PDM / I2S / TDM"),
        ("② VAD", "μW 级，模拟或极简数字，always-on"),
        ("③ KWS Conv1D", "<1 mW，检测到活动话语时运行"),
        ("④ 命中唤醒词", "中断 PMU → 上电 NPU / CPU 域"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table, 16.5)
    for j, (title, detail) in enumerate(steps):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, "E8EEF5" if j % 2 == 0 else "F2F5F8")
        set_cell_borders(cell, color="2E5A88", sz="8")
        set_cell_margins(cell, top=60, bottom=60, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, space_before=0, space_after=2)
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p2, space_before=0, space_after=0)
        run = p2.add_run(detail)
        set_run_font(run, size=8, color=GRAY)

    # arrow hint
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=2, space_after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("全部位于 Audio Core 独立低功耗电源域（自有时钟 + retention SRAM）")
    set_run_font(run, size=9, italic=True, color=GRAY)
    return table


def parse_table_block(lines: list[str], start: int):
    """Parse a markdown table starting at start. Return (headers, rows, next_index)."""
    header_line = lines[start].strip()
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    # skip separator
    i = start + 1
    if i < len(lines) and re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i].strip()):
        i += 1
    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.strip("|").split("|")]
        # pad/truncate
        if len(cells) < len(headers):
            cells += [""] * (len(headers) - len(cells))
        rows.append(cells[: len(headers)])
        i += 1
    return headers, rows, i


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    rPr = normal._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")


def convert(md_text: str, doc: Document):
    lines = md_text.splitlines()
    i = 0
    title_done = False
    subtitle_done = False
    in_code = False
    code_lines: list[str] = []
    toc_items: list[str] = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip("\n")
        stripped = raw.strip()

        # code fence
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                # Detect special diagrams and replace with clean layouts
                joined = "\n".join(code_lines)
                if "Audio Core" in joined and "AFE Engine" in joined:
                    add_architecture_diagram(doc)
                elif "PDM mic" in joined and "always-on" in joined:
                    add_wakeup_chain(doc)
                else:
                    add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # empty
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1
            continue

        # title (# )
        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_heading(doc, stripped[2:], level=0)
            title_done = True
            i += 1
            continue

        # subtitle-ish ## under title before TOC
        if stripped.startswith("## ") and not title_done:
            add_heading(doc, stripped[3:], level=1)
            i += 1
            continue

        # special: first ## as subtitle (after title)
        if stripped.startswith("## ——") or (stripped.startswith("## ") and not subtitle_done and title_done and "目录" not in stripped):
            # treat decorative subtitle as centered gray
            p = doc.add_paragraph()
            set_paragraph_format(p, space_before=0, space_after=10, line_spacing=1.2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[3:].strip())
            set_run_font(run, size=12, italic=True, color=GRAY)
            subtitle_done = True
            i += 1
            continue

        # headings
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if text == "目录":
                add_heading(doc, text, level=1)
                # collect TOC from following numbered list until blank/next heading
                i += 1
                while i < len(lines):
                    t = lines[i].strip()
                    if not t:
                        # peek ahead; if next non-empty is not numbered, stop
                        j = i + 1
                        while j < len(lines) and not lines[j].strip():
                            j += 1
                        if j >= len(lines) or not re.match(r"^\d+\.", lines[j].strip()):
                            break
                        i += 1
                        continue
                    if t.startswith("#") or t.startswith("---"):
                        break
                    m = re.match(r"^(\d+)\.\s+\[([^\]]+)\]", t)
                    if m:
                        toc_items.append(f"{m.group(1)}. {m.group(2)}")
                        add_body(doc, f"{m.group(1)}. {m.group(2)}")
                    elif re.match(r"^\d+\.", t):
                        add_body(doc, re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t))
                    else:
                        break
                    i += 1
                continue
            add_heading(doc, text, level=1)
            i += 1
            continue

        # blockquote / note
        if stripped.startswith(">"):
            note_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                note_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            add_note(doc, note_lines)
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and re.search(r"\|?\s*:?-+:?\s*\|", lines[i + 1]):
            headers, rows, next_i = parse_table_block(lines, i)
            add_table(doc, headers, rows)
            i = next_i
            continue

        # numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            add_numbered(doc, m.group(2), int(m.group(1)))
            i += 1
            continue

        # nested bullet with spaces
        m = re.match(r"^(\s*)[-*]\s+(.*)$", raw)
        if m:
            indent_spaces = len(m.group(1).replace("\t", "    "))
            level = 0 if indent_spaces < 2 else 1
            add_bullet(doc, m.group(2), level=level)
            i += 1
            continue

        # plain paragraph — strip leftover markdown links
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", stripped)
        add_body(doc, text)
        i += 1

    # footer note
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=18, space_after=0, line_spacing=1.2)
    run = p.add_run("— 报告完 —")
    set_run_font(run, size=10, italic=True, color=GRAY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    convert(md, doc)
    doc.save(DST)
    UPLOAD.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DST, UPLOAD)
    print(f"Wrote {DST} ({DST.stat().st_size} bytes)")
    print(f"Copied {UPLOAD}")


if __name__ == "__main__":
    main()
