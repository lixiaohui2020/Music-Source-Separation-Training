#!/usr/bin/env python3
"""Convert the revised proposal markdown to a styled .docx file."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

MD_PATH = Path("docs/algorithm-automation-reports/08-算法自动化工具流程修订版方案.md")
OUT_PATH = Path("docs/algorithm-automation-reports/08-算法自动化工具流程修订版方案.docx")

HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),
    2: RGBColor(0x1E, 0x40, 0xAF),
    3: RGBColor(0x33, 0x41, 0x55),
    4: RGBColor(0x47, 0x55, 0x69),
}


def set_zh_font(run, name: str = "微软雅黑") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=min(level, 4))
    run = p.add_run(strip_inline(text))
    run.font.color.rgb = HEADING_COLORS.get(min(level, 4), HEADING_COLORS[4])
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    run.font.size = Pt(sizes.get(min(level, 4), 12))
    run.font.bold = True
    set_zh_font(run)


def add_paragraph(doc: Document, text: str, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(strip_inline(part))
            run.font.bold = True
        else:
            run = p.add_run(strip_inline(part))
            if bold:
                run.font.bold = True
        run.font.size = Pt(10.5)
        set_zh_font(run)


def add_list_item(doc: Document, text: str, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        text = ("• " if not ordered else "") + text
    for part in re.split(r"(\*\*.+?\*\*)", text):
        run = p.add_run(strip_inline(part))
        if part.startswith("**"):
            run.font.bold = True
        run.font.size = Pt(10.5)
        set_zh_font(run)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    run = p.add_run(strip_inline(text))
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    set_zh_font(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", line.strip()):
            continue
        rows.append([strip_inline(c.strip()) for c in line.strip().strip("|").split("|")])
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(row[ci] if ci < len(row) else "")
            run.font.size = Pt(9)
            if ri == 0:
                run.font.bold = True
            set_zh_font(run)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, parse_table(table_buf))
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
        elif stripped == "---":
            pass
        elif stripped.startswith("> "):
            add_quote(doc, stripped[2:])
        elif re.match(r"^[-*]\s+", stripped):
            add_list_item(doc, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            add_list_item(doc, re.sub(r"^\d+[.)]\s+", "", stripped), ordered=True)
        else:
            add_paragraph(doc, stripped)
        i += 1

    if in_code and code_buf:
        add_code_block(doc, code_buf)
    flush_table()

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    if not MD_PATH.exists():
        print(f"Markdown not found: {MD_PATH}", file=sys.stderr)
        raise SystemExit(1)
    convert(MD_PATH, OUT_PATH)
